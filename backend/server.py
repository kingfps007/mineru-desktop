# -*- coding: utf-8 -*-
"""
MinerU Desktop v3.0.4 — Python Backend API Server
FastAPI server providing REST APIs for GPU detection, setup wizard,
Zotero JSON parsing, batch PDF processing, Markdown-to-Word generation,
and MinerU Cloud API integration.
"""
import os, sys, json, re, time, subprocess, threading, uuid, shutil, asyncio, gc, io, zipfile
from pathlib import Path
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
import urllib.request, urllib.error
import uvicorn

# ── Constants ──
PROJECT_ROOT = Path(__file__).resolve().parent.parent
CONFIG_PATH = Path.home() / "mineru_desktop_config.json"
DEFAULT_PORT = 18766
MAX_PARSE_TIMEOUT = 600
GPU_TEMP_PAUSE_THRESHOLD = 70
GPU_TEMP_PAUSE_SECONDS = 30
BATCH_SIZE = 10
MAX_RETRIES = 3

executor = ThreadPoolExecutor(max_workers=4)

# ── In-memory state ──
state = {
    "setup_running": False,
    "parse_running": False,
    "parse_cancel_requested": False,
    "parse_paused": False,
    "parse_pause_reason": "",
    "parse_progress": {"current": 0, "total": 0, "current_file": "", "file_progress": 0},
    "parse_history": {},
    "docx_running": False,
    "imported_references": [],
    "direct_files": [],
    "log_clients": set(),
    "_current_proc": None,
    "_current_stop": None,
}

# ── GPU / CPU / Process Cleanup ──
def cleanup_memory():
    released = []
    # 1. Kill zombie mineru child processes (Windows: taskkill /F, NOT SIGTERM)
    killed = 0
    try:
        r = subprocess.run(['tasklist', '/fo', 'csv', '/nh'], capture_output=True, text=True, timeout=5)
        for line in r.stdout.splitlines():
            low = line.lower()
            if 'mineru-api' in low or 'mineru' in low and '.exe' in low:
                pid_str = line.split(',')[1].strip('"')
                if pid_str.isdigit():
                    # Windows 强制终止（taskkill 比 os.kill 可靠，避免卡死）
                    subprocess.run(['taskkill', '/F', '/PID', pid_str], capture_output=True, timeout=3)
                    killed += 1
    except Exception as e:
        sync_log(f"⚠ 清理进程时出错: {e}")
    if killed: released.append(f"终止 {killed} 个僵尸进程")

    # 2. Python GC
    before = _get_mem_mb()
    gc.collect()
    after = _get_mem_mb()
    if before - after > 10: released.append(f"释放内存 {before - after:.0f} MB")

    # 3. GPU VRAM cleanup via torch (动态路径，超时保护)
    try:
        env = cached_detect_mineru_env()
        python_exe = env["python"] if env and env.get("python") else "python"
        subprocess.run(
            [python_exe, "-c",
             "import gc,torch; gc.collect(); torch.cuda.empty_cache(); print(torch.cuda.memory_allocated()//1024**2)"],
            capture_output=True, text=True, timeout=10)
    except subprocess.TimeoutExpired:
        pass
    except Exception:
        pass

    if not released: released.append("内存已是最优状态")
    msg = "🧹 " + " | ".join(released)
    sync_log(msg)
    return {"message": msg, "killed": killed}

def _get_mem_mb():
    try:
        import ctypes
        class MEMORYSTATUSEX(ctypes.Structure):
            _fields_ = [("dwLength", ctypes.c_ulong), ("dwMemoryLoad", ctypes.c_ulong),
                         ("ullTotalPhys", ctypes.c_ulonglong), ("ullAvailPhys", ctypes.c_ulonglong),
                         ("ullTotalPageFile", ctypes.c_ulonglong), ("ullAvailPageFile", ctypes.c_ulonglong),
                         ("ullTotalVirtual", ctypes.c_ulonglong), ("ullAvailVirtual", ctypes.c_ulonglong)]
        m = MEMORYSTATUSEX(); m.dwLength = ctypes.sizeof(MEMORYSTATUSEX)
        ctypes.windll.kernel32.GlobalMemoryStatusEx(ctypes.byref(m))
        return (m.ullTotalPhys - m.ullAvailPhys) / (1024**2)
    except: return 0

# ═══════════════════════════════════════════════════════════
# 1. Detection Functions
# ═══════════════════════════════════════════════════════════

def find_conda():
    paths = [
        Path("C:/ProgramData/miniconda3/Scripts/conda.exe"),
        Path("C:/ProgramData/anaconda3/Scripts/conda.exe"),
        Path.home() / "miniconda3/Scripts/conda.exe",
        Path.home() / "anaconda3/Scripts/conda.exe",
    ]
    for p in paths:
        if p.exists():
            return str(p)
    try:
        r = subprocess.run(["where", "conda"], capture_output=True, text=True, timeout=10)
        for line in r.stdout.strip().split("\n"):
            line = line.strip()
            if line.lower().endswith("conda.exe"):
                return line
    except:
        pass
    return None

def detect_mineru_env():
    candidates = []
    env_dirs = [
        Path.home() / ".conda" / "envs",
        Path("C:/ProgramData/miniconda3/envs"),
        Path("C:/ProgramData/anaconda3/envs"),
    ]
    for base in env_dirs:
        env = base / "MinerU"
        if (env / "python.exe").exists():
            candidates.append(env)
    users = Path("C:/Users")
    if users.exists():
        for ud in users.iterdir():
            env = ud / ".conda" / "envs" / "MinerU"
            if (env / "python.exe").exists() and env not in candidates:
                candidates.append(env)
    if not candidates:
        return None
    env = candidates[0]
    mineru_exe = env / "Scripts" / "mineru.exe"
    dl_exe = env / "Scripts" / "mineru-models-download.exe"
    return {
        "root": str(env), "python": str(env / "python.exe"),
        "mineru": str(mineru_exe) if mineru_exe.exists() else None,
        "mineru_dl": str(dl_exe) if dl_exe.exists() else None,
        "version": "3.1.x", "ok": True,
    }

def detect_gpu_fast():
    """Fast GPU detection: nvidia-smi only. Returns <50ms."""
    info = {"gpu_name": "N/A", "vram_mb": 0, "driver": "N/A", "temperature": 0,
            "has_nvidia": False, "cuda_avail": None, "config_cuda": False}
    try:
        r = subprocess.run(
            ["nvidia-smi", "--query-gpu=name,memory.total,driver_version,temperature.gpu",
             "--format=csv,noheader"], capture_output=True, text=True, timeout=5)
        if r.returncode == 0 and r.stdout.strip():
            parts = [x.strip() for x in r.stdout.strip().split(",")]
            if len(parts) >= 1: info["gpu_name"] = parts[0]
            if len(parts) >= 2: info["vram_mb"] = int(parts[1].replace(" MiB", "").replace(" ", ""))
            if len(parts) >= 3: info["driver"] = parts[2]
            if len(parts) >= 4:
                try: info["temperature"] = int(parts[3])
                except: pass
            info["has_nvidia"] = True
    except: pass
    config_path = Path.home() / "mineru.json"
    if config_path.exists():
        try:
            cfg = json.loads(config_path.read_text(encoding="utf-8"))
            info["config_cuda"] = cfg.get("device-mode") == "cuda"
        except: pass
    return info

def detect_gpu_deep():
    """Deep GPU detection: includes torch.cuda.is_available()."""
    info = dict(detect_gpu_fast())
    env = cached_detect_mineru_env()
    if env and info["has_nvidia"]:
        try:
            r = subprocess.run(
                [env["python"], "-c", "import torch; print(torch.cuda.is_available())"],
                capture_output=True, text=True, timeout=30)
            if r.returncode == 0:
                info["cuda_avail"] = r.stdout.strip() == "True"
        except: pass
    return info

def detect_gpu():
    return detect_gpu_fast()

def detect_models_fast():
    ms = Path.home() / ".cache" / "modelscope" / "hub" / "models" / "OpenDataLab"
    hf = Path.home() / ".cache" / "huggingface" / "hub"
    result = {"pipeline": False, "vlm": False, "pipeline_mb": 0, "vlm_mb": 0}

    config_path = Path.home() / "mineru.json"
    if config_path.exists():
        try:
            cfg = json.loads(config_path.read_text(encoding="utf-8"))
            pp = cfg.get("models-dir", {}).get("pipeline", "")
            vp = cfg.get("models-dir", {}).get("vlm", "")
            if pp and Path(pp).exists(): result["pipeline"] = True
            if vp and Path(vp).exists(): result["vlm"] = True
        except: pass

    if not result["pipeline"]:
        pipe_key = ms / "PDF-Extract-Kit-1___0" / "models" / "Layout" / "PP-DocLayoutV2" / "model.safetensors"
        if pipe_key.exists(): result["pipeline"] = True

    if not result["vlm"] and ms.exists():
        for p in [ms / "MinerU2___5-Pro-2604-1___2B" / "model.safetensors"]:
            if p.exists(): result["vlm"] = True; break
        if not result["vlm"]:
            for d in ms.iterdir():
                dn = d.name.lower()
                if any(k in dn for k in ['mineru2', 'vlm', '1___2b']):
                    if (d / "model.safetensors").exists(): result["vlm"] = True; break

    if not result["pipeline"]:
        for snap in hf.glob("models--opendatalab--PDF-Extract-Kit-1.0/snapshots/*"):
            if snap.is_dir() and any(snap.iterdir()): result["pipeline"] = True; break
    if not result["vlm"]:
        for snap in hf.glob("models--opendatalab--MinerU2.5-Pro-2604-1.2B/snapshots/*"):
            if snap.is_dir() and any(snap.iterdir()): result["vlm"] = True; break
    return result

def detect_models():
    return detect_models_fast()

# ── Detection cache ──
_detection_cache = {}

def _cached_detect(key, fn, ttl=5):
    now = time.time()
    entry = _detection_cache.get(key)
    if entry and now - entry["ts"] < ttl:
        return entry["data"]
    data = fn()
    _detection_cache[key] = {"ts": now, "data": data}
    return data

def cached_detect_models():
    return _cached_detect("models", detect_models_fast, ttl=3)

def cached_detect_gpu():
    return _cached_detect("gpu", detect_gpu_fast, ttl=3)

def cached_find_conda():
    return _cached_detect("conda", find_conda, ttl=10)

def cached_detect_mineru_env():
    return _cached_detect("mineru_env", detect_mineru_env, ttl=5)

def detect_system_ram():
    try:
        import ctypes
        class MEMORYSTATUSEX(ctypes.Structure):
            _fields_ = [("dwLength", ctypes.c_ulong), ("dwMemoryLoad", ctypes.c_ulong),
                         ("ullTotalPhys", ctypes.c_ulonglong), ("ullAvailPhys", ctypes.c_ulonglong),
                         ("ullTotalPageFile", ctypes.c_ulonglong), ("ullAvailPageFile", ctypes.c_ulonglong),
                         ("ullTotalVirtual", ctypes.c_ulonglong), ("ullAvailVirtual", ctypes.c_ulonglong)]
        m = MEMORYSTATUSEX()
        m.dwLength = ctypes.sizeof(MEMORYSTATUSEX)
        ctypes.windll.kernel32.GlobalMemoryStatusEx(ctypes.byref(m))
        return round(m.ullTotalPhys / (1024**3), 1)
    except: return 0

def capability_assessment():
    gpu = cached_detect_gpu()
    models = cached_detect_models()
    env = cached_detect_mineru_env()
    conda = cached_find_conda()
    has_nvidia = gpu["has_nvidia"]
    vram_gb = gpu["vram_mb"] / 1024 if gpu["vram_mb"] else 0

    # 综合评估：必须 conda + mineru环境 + 模型齐全 才算环境就绪
    has_conda = bool(conda)
    has_mineru = env is not None
    has_models = models.get("pipeline", False) and models.get("vlm", False)

    env_ready = has_conda and has_mineru and has_models

    modes = []
    if env and models["pipeline"]: modes.append("pipeline")
    if env and models["vlm"] and has_nvidia: modes.append("vlm")

    # level 决定 UI 显示和导航启用：先看 env_ready
    if not env_ready:
        # 环境未就绪，强制 basic，提示走安装向导
        missing = []
        if not has_conda: missing.append("Miniconda")
        if not has_mineru: missing.append("MinerU 环境")
        if not has_models: missing.append("模型")
        recommendation = f"环境未就绪，缺少: {', '.join(missing)}。请先完成安装向导。"
        level = "basic"
    elif not has_nvidia:
        recommendation = "未检测到 NVIDIA 显卡。建议使用 Pipeline (传统规则) 模式。"
        level = "limited"
    elif vram_gb < 4:
        recommendation = f"检测到 {gpu['gpu_name']} ({vram_gb:.1f}GB 显存)。显存不足以运行 VLM 模型，仅支持 Pipeline 模式。"
        level = "limited"
    elif vram_gb < 6:
        recommendation = f"检测到 {gpu['gpu_name']} ({vram_gb:.1f}GB 显存)。VLM 可用但显存紧张。"
        level = "capable"
    else:
        recommendation = f"检测到 {gpu['gpu_name']} ({vram_gb:.1f}GB 显存)。完全支持 VLM + Pipeline。"
        level = "full"

    return {
        "has_nvidia": has_nvidia, "vram_gb": round(vram_gb, 1),
        "available_modes": modes, "pipeline_available": "pipeline" in modes,
        "vlm_available": "vlm" in modes,
        "level": level, "recommendation": recommendation,
        "conda_installed": has_conda, "mineru_installed": has_mineru,
        "models_ready": has_models, "env_ready": env_ready,
    }

def get_gpu_temperature():
    try:
        r = subprocess.run(
            ["nvidia-smi", "--query-gpu=temperature.gpu", "--format=csv,noheader"],
            capture_output=True, text=True, timeout=5)
        if r.returncode == 0 and r.stdout.strip():
            return int(r.stdout.strip())
    except: pass
    return 0

def load_config():
    if CONFIG_PATH.exists():
        try: return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
        except: return {}
    return {}

def save_config(config):
    CONFIG_PATH.write_text(json.dumps(config, indent=2, ensure_ascii=False), encoding="utf-8")

# ═══════════════════════════════════════════════════════════
# 2. Zotero JSON Parsing
# ═══════════════════════════════════════════════════════════

def parse_zotero_json(file_path: str) -> list:
    with open(file_path, "r", encoding="utf-8-sig") as f:
        data = json.load(f)
    items = data if isinstance(data, list) else data.get("items", [])
    references = []
    for item in items:
        citekey = item.get("citationKey", "").strip()
        title = item.get("title", "")
        uri = item.get("uri", "")
        pdfs = []
        for att in item.get("attachments", []) or []:
            path = att.get("path", "")
            if not path:
                sp = att.get("storedPaths")
                if sp and isinstance(sp, list) and len(sp) > 0:
                    path = sp[0] or ""
            if path and path.lower().endswith(".pdf"):
                pdfs.append(path)
        pdf_exists = False
        pdf_path = ""
        pdf_size = 0
        for p in pdfs:
            full_path = Path(p)
            if full_path.exists():
                pdf_exists = True; pdf_path = str(full_path); pdf_size = full_path.stat().st_size; break
        if not pdf_exists and pdfs: pdf_path = pdfs[0]
        references.append({
            "citekey": citekey, "title": title, "uri": uri,
            "pdf_exists": pdf_exists, "pdf_path": pdf_path, "pdf_size": pdf_size,
            "item_type": item.get("itemType", ""), "date": item.get("date", ""),
            "creators": item.get("creators", []), "publicationTitle": item.get("publicationTitle", ""),
            "volume": item.get("volume", ""), "pages": item.get("pages", ""),
            "DOI": item.get("DOI", ""), "ISSN": item.get("ISSN", ""),
            "journalAbbreviation": item.get("journalAbbreviation", ""),
        })
    return references

# ═══════════════════════════════════════════════════════════
# 3. Setup Wizard Backend
# ═══════════════════════════════════════════════════════════

SETUP_STEPS_DEF = [
    {"id": "check_env", "name": "环境扫描", "description": "自动检测系统中已安装的开发工具，包括 Miniconda/Anaconda、NVIDIA 驱动、CUDA 支持、MinerU 环境及 AI 模型。", "location": "扫描 C:/ProgramData、用户 .conda 目录和 NVIDIA 驱动", "size": "无需额外磁盘空间", "action_text": "开始扫描", "check_done": lambda: bool(cached_find_conda())},
    {"id": "install_conda", "name": "安装 Miniconda", "description": "Miniconda 是一个轻量级 Python 环境管理器，MinerU 依赖它来创建独立的 Python 运行环境。未安装时将打开国内镜像下载页。", "location": "默认安装到 C:/ProgramData/miniconda3 (约 400 MB)", "size": "约 400 MB", "action_text": "打开下载页面", "check_done": lambda: bool(cached_find_conda())},
    {"id": "create_env", "name": "创建 MinerU 环境", "description": "使用 Conda 创建一个 Python 3.10 虚拟环境 'MinerU'，后续所有 MinerU 的依赖包都将安装在这个隔离环境中。", "location": "C:/ProgramData/miniconda3/envs/MinerU 或 C:/Users/<用户>/.conda/envs/MinerU", "size": "基础环境约 150 MB，安装依赖后约 3-5 GB", "action_text": "创建环境", "check_done": lambda: cached_detect_mineru_env() is not None},
    {"id": "install_torch", "name": "安装 PyTorch CUDA", "description": "PyTorch 是 MinerU 的深度学习框架后端。CUDA 版本利用 NVIDIA GPU 加速解析，速度可提升 10-50 倍。无 NVIDIA 显卡将自动跳过此步骤。", "location": "安装到 Conda MinerU 环境: <MinerU 环境>", "size": "约 2.5 GB (含 torch + torchvision + CUDA 12.1)", "action_text": "安装 PyTorch", "check_done": lambda: _check_torch_installed()},
    {"id": "install_mineru", "name": "安装 MinerU 程序包", "description": "pip 安装 MinerU PDF 解析引擎 + hf_xet 加速下载。MinerU 支持 Pipeline（传统规则引擎）和 VLM（视觉大模型）两种后端，自动处理数学公式、表格和图片提取。", "location": "安装到 Conda MinerU 环境的 site-packages 目录", "size": "约 500 MB (含 transformers + hf_xet 等依赖)", "action_text": "安装 MinerU", "check_done": lambda: cached_detect_mineru_env() is not None},
    {"id": "download_pipeline", "name": "下载 Pipeline 模型", "description": "下载 Pipeline 模型（约 4GB），用于布局分析、表格识别和公式提取。这是基本模型，所有用户都建议下载，CPU 也可使用。", "location": "C:/Users/<用户>/.cache/modelscope/hub/models/OpenDataLab/", "size": "约 4 GB", "action_text": "下载 Pipeline 模型", "check_done": lambda: cached_detect_models()["pipeline"]},
    {"id": "download_vlm", "name": "下载 VLM 模型", "description": "下载 VLM 视觉大模型（约 2.2GB），仅 NVIDIA GPU 用户需要。无英伟达显卡请跳过此步，使用云端 API 替代。", "location": "C:/Users/<用户>/.cache/modelscope/hub/models/OpenDataLab/", "size": "约 2.2 GB", "action_text": "下载 VLM 模型", "check_done": lambda: cached_detect_models()["vlm"], "skip_if": lambda: not detect_gpu_fast()["has_nvidia"]},
    {"id": "install_zotero", "name": "安装 Zotero", "description": "Zotero 是开源文献管理软件 (最新 Zotero 9，~227 MB)，用于管理论文 PDF 和引用信息。此步骤仅打开下载页面，需手动安装。", "location": "C:/Program Files/Zotero (默认安装路径)", "size": "约 227 MB (Zotero 9)", "action_text": "打开下载页", "check_done": lambda: True},
    {"id": "install_bbt", "name": "安装 Zotero 插件", "description": "推荐从 Zotero 中文插件商店下载 Add-on Market (插件市场) 和 Better BibTeX v9.0.27 (~32.6 MB)。安装后重启 Zotero，右键文献库 → Export → BetterBibTeX JSON。", "location": "Zotero 内部 plugins 目录", "size": "BBT 约 32.6 MB / 插件市场约 2 MB", "action_text": "打开中文插件商店", "check_done": lambda: True},
    {"id": "enable_cuda", "name": "启用 GPU 加速", "description": "修改 mineru.json 配置文件，将 device-mode 设为 'cuda'，使 MinerU 默认使用 NVIDIA GPU 进行推理加速。", "location": "C:/Users/<用户>/mineru.json", "size": "无需额外磁盘空间", "action_text": "启用 CUDA", "check_done": lambda: _check_cuda_enabled(), "skip_if": lambda: not detect_gpu_fast()["has_nvidia"]},
]

def _check_torch_installed():
    env = cached_detect_mineru_env()
    if not env: return False
    try:
        r = subprocess.run([env["python"], "-c", "import torch; print(torch.cuda.is_available())"],
                           capture_output=True, text=True, timeout=20)
        return r.returncode == 0
    except: return False

def _check_zotero_installed():
    zotero_exe = Path("C:/Program Files/Zotero/Zotero.exe")
    return zotero_exe.exists()

def _check_models_done():
    models = cached_detect_models()
    return models["pipeline"] and models["vlm"]

def _check_cuda_enabled():
    config_path = Path.home() / "mineru.json"
    if config_path.exists():
        try:
            cfg = json.loads(config_path.read_text(encoding="utf-8"))
            return cfg.get("device-mode") == "cuda"
        except: pass
    return False

# ── SetupRunner ──
class SetupRunner:
    def __init__(self, log_cb, progress_cb):
        self.log = log_cb; self.progress = progress_cb
        self._conda_path = cached_find_conda()

    def _conda(self, args, timeout=600):
        if not self._conda_path: self._conda_path = cached_find_conda()
        if not self._conda_path: self.log("✗ 找不到 conda.exe"); return 1
        return self._run([self._conda_path] + args, timeout)

    def _pip(self, args, timeout=600):
        if not self._conda_path: self._conda_path = cached_find_conda()
        return self._run([self._conda_path, "run", "-n", "MinerU", "pip"] + args, timeout)

    def _mineru_exe(self, args, timeout=3600):
        env = cached_detect_mineru_env()
        if not env: self.log("✗ MinerU 环境未安装"); return 1
        cmd = [env["mineru_dl"]] + args if args[0] != "--version" else [env["mineru"]] + args
        return self._run(cmd, timeout)

    def _run(self, cmd, timeout=3600):
        env = os.environ.copy()
        env["HF_ENDPOINT"] = "https://hf-mirror.com"
        proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                                text=True, encoding="utf-8", errors="replace", env=env)
        last_pct = -1
        for line in proc.stdout:
            line = line.strip()
            if line: self.log(line)
            m = re.search(r"(\d+)%", line)
            if m:
                pct = int(m.group(1))
                if pct != last_pct: last_pct = pct; self.progress(pct)
        proc.wait(timeout=timeout)
        return proc.returncode

    def execute_step(self, step_id):
        _detection_cache.clear()
        steps = {
            "check_env": ("step_check_env", self._step_check_env),
            "install_conda": ("step_install_conda", self._step_install_conda),
            "create_env": ("step_create_env", self._step_create_env),
            "install_torch": ("step_install_torch", self._step_install_torch),
            "install_mineru": ("step_install_mineru", self._step_install_mineru),
            "download_pipeline": ("step_download_pipeline", self._step_download_pipeline),
            "download_vlm": ("step_download_vlm", self._step_download_vlm),
            "install_zotero": ("step_install_zotero", self._step_install_zotero),
            "install_bbt": ("step_install_bbt", self._step_install_bbt),
            "enable_cuda": ("step_enable_cuda", self._step_enable_cuda),
        }
        if step_id not in steps: self.log(f"✗ 未知步骤: {step_id}"); return False
        step_name, func = steps[step_id]
        self.log(f"\n── {step_name} ──")
        self.progress(0)
        try:
            result = func()
            self.progress(100 if result == 0 else 0)
            return result == 0
        except Exception as e:
            self.log(f"✗ {step_name} 失败: {e}"); self.progress(0); return False

    def _step_check_env(self):
        conda = cached_find_conda(); gpu = cached_detect_gpu()
        env = cached_detect_mineru_env(); models = cached_detect_models()
        self.log(f"{'✓' if conda else '✗'} Conda: {'已安装' if conda else '未安装'}")
        if gpu["has_nvidia"]:
            self.log(f"✓ GPU: {gpu['gpu_name']} ({gpu['vram_mb']} MB, 驱动 {gpu['driver']})")
            if gpu["vram_mb"] < 4096: self.log("  ⚠ 显存<4GB，不建议下载 VLM 模型")
        else: self.log("⚠ GPU: 未检测到 NVIDIA 显卡")
        self.log(f"{'✓' if env else '✗'} MinerU: {'已安装' if env else '未安装'}")
        self.log(f"{'✓' if models['pipeline'] else '✗'} Pipeline: {'已下载' if models['pipeline'] else '未下载'}")
        self.log(f"{'✓' if models['vlm'] else '✗'} VLM: {'已下载' if models['vlm'] else '未下载'}")
        return 0

    def _step_install_conda(self):
        p = cached_find_conda()
        if p: self.log(f"✓ Conda 已安装: {p}"); self._conda_path = p; return 0
        self.log("未检测到 Miniconda，正在打开下载页面...")
        self.log("  https://mirrors.tuna.tsinghua.edu.cn/anaconda/miniconda/")
        self.log("安装时选择「Just Me」→ 全部默认即可")
        import webbrowser; webbrowser.open("https://mirrors.tuna.tsinghua.edu.cn/anaconda/miniconda/")
        self.log("⚠ 请手动安装 Miniconda 后点击「重新检测」"); return 1

    def _step_create_env(self):
        env = cached_detect_mineru_env()
        if env: self.log(f"✓ MinerU 环境已存在: {env['root']}"); return 0
        self.log("创建 conda 环境 MinerU (Python 3.10)..."); return self._conda(["create", "-n", "MinerU", "python=3.10", "-y"], timeout=600)

    def _step_install_torch(self):
        gpu = cached_detect_gpu()
        if not gpu["has_nvidia"]: self.log("⚠ 跳过 (无 NVIDIA GPU，建议使用云端 API)"); return 0
        if _check_torch_installed(): self.log("✓ PyTorch CUDA 已就绪"); return 0
        self.log("安装 PyTorch CUDA 12.1 ..."); return self._pip(["install", "torch", "torchvision", "--index-url", "https://download.pytorch.org/whl/cu121"], timeout=1200)

    def _step_install_mineru(self):
        self.log("安装 MinerU + hf_xet ..."); ret = self._pip(["install", "mineru"], timeout=600)
        if ret == 0: self.log("安装 hf_xet (HuggingFace 加速传输)..."); self._pip(["install", "hf_xet"], timeout=120)
        return ret

    def _step_download_pipeline(self):
        models = cached_detect_models()
        if models["pipeline"]: self.log("✓ Pipeline 模型已下载"); return 0
        self.log("下载 Pipeline 模型 (~4GB)..."); return self._mineru_exe(["-s", "modelscope", "-m", "pipeline"])

    def _step_download_vlm(self):
        gpu = cached_detect_gpu()
        if not gpu["has_nvidia"]: self.log("⚠ 跳过 (无 NVIDIA GPU，VLM 模型无法使用)"); return 0
        if gpu["vram_mb"] < 4096: self.log("⚠ 跳过 (显存<4GB，VLM 模型无法运行)"); return 0
        models = cached_detect_models()
        if models["vlm"]: self.log("✓ VLM 模型已下载"); return 0
        self.log("下载 VLM 模型 (~2.2GB)..."); return self._mineru_exe(["-s", "modelscope", "-m", "vlm"])

    def _step_install_zotero(self):
        self.log("打开 Zotero 下载页面: https://www.zotero.org/download/")
        self.log("Zotero 9 安装包约 227 MB，安装后创建账户并登录")
        import webbrowser; webbrowser.open("https://www.zotero.org/download/")
        self.log("⚠ 请手动安装 Zotero")
        return 1

    def _step_install_bbt(self):
        self.log("打开 Zotero 中文插件商店: https://zotero-chinese.com/plugins/")
        self.log("推荐下载两个插件:")
        self.log("  1. Add-on Market for Zotero (插件市场)")
        self.log("  2. Better BibTeX for Zotero v9.0.27 (~32.6 MB)")
        self.log("下载 .xpi 后: Zotero → 工具 → 插件 → Install from File → 选择 .xpi → 重启")
        self.log("Zotero 7+ 也可直接在插件市场搜索 'BetterBibTeX' 安装")
        import webbrowser; webbrowser.open("https://zotero-chinese.com/plugins/")
        self.log("⚠ 请手动安装插件后重启 Zotero")
        return 1

    def _step_enable_cuda(self):
        gpu = cached_detect_gpu()
        if not gpu["has_nvidia"]: self.log("⚠ 跳过 (无 NVIDIA GPU)"); return 0
        config_path = Path.home() / "mineru.json"
        if config_path.exists():
            try:
                c = json.loads(config_path.read_text(encoding="utf-8"))
                c["device-mode"] = "cuda"; config_path.write_text(json.dumps(c, indent=2, ensure_ascii=False), encoding="utf-8")
                self.log("✓ mineru.json: device-mode = cuda"); return 0
            except: self.log("⚠ 无法修改配置文件"); return 1
        self.log("⚠ mineru.json 不存在（模型下载后自动生成）"); return 1

# ═══════════════════════════════════════════════════════════
# 4. Build DOCX
# ═══════════════════════════════════════════════════════════

def build_docx_with_zotero(md_path: str, json_path: str, output_dir: str, log_cb):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def load_bbt():
        with open(json_path, "r", encoding="utf-8-sig") as f: data = json.loads(f.read())
        key2uri = {}; key2info = {}
        for item in data.get("items", []):
            ck = item.get("citationKey", "").strip(); uri = item.get("uri", "").strip()
            if not ck or not uri: continue
            creators = item.get("creators", [])
            authors = [{"family": c.get("lastName", ""), "given": c.get("firstName", "")} for c in creators if c.get("creatorType") == "author"]
            date_str = item.get("date", ""); date_parts = [[]]
            if date_str:
                parts = re.split(r'[/-]', date_str)
                for p in parts:
                    try: date_parts[0].append(int(p.strip()))
                    except: pass
            item_data = {"id": item.get("itemID", 0), "type": "article-journal", "title": item.get("title", ""), "author": authors}
            if date_parts[0]: item_data["issued"] = {"date-parts": [date_parts[0]]}
            if item.get("publicationTitle"): item_data["container-title"] = item["publicationTitle"]
            if item.get("volume"): item_data["volume"] = item["volume"]
            if item.get("pages"): item_data["page"] = item["pages"]
            if item.get("DOI"): item_data["DOI"] = item["DOI"]
            if item.get("ISSN"): item_data["ISSN"] = item["ISSN"]
            if item.get("journalAbbreviation"): item_data["container-title-short"] = item["journalAbbreviation"]
            key2uri[ck] = uri; key2info[ck] = item_data
        return key2uri, key2info

    def make_citation_json(keys, key2uri, key2info):
        items = []; first_author = ""
        for ck in keys:
            if ck not in key2uri: continue
            cit_item = {"uris": [key2uri[ck]], "itemData": key2info.get(ck, {"id": 0, "type": "article-journal", "title": ck})}
            items.append(cit_item)
            if not first_author:
                auths = key2info.get(ck, {}).get("author", [])
                first_author = auths[0].get("family", ck) if auths else ck
        if not items: return None
        author_label = first_author if first_author else ", ".join(keys[:2])
        return json.dumps({"citationID": str(uuid.uuid4()), "properties": {"formattedCitation": f"({author_label})", "plainCitation": f"({author_label})", "noteIndex": 0}, "citationItems": items}, ensure_ascii=False)

    def add_field_code(paragraph, field_json_str):
        run = paragraph.add_run(" ")
        for (cht, ct) in [('begin', 'w:fldChar'), ('instrText', 'w:instrText'), ('separate', 'w:fldChar')]:
            el = OxmlElement(ct)
            if cht == 'fldChar': el.set(qn('w:fldCharType'), 'begin' if ct.endswith('fldChar') else 'separate')
            elif cht == 'instrText':
                el.set(qn('xml:space'), 'preserve')
                el.text = f' ADDIN ZOTERO_ITEM CSL_CITATION {field_json_str}'
            run._r.append(el)
        run._r.append(OxmlElement('w:fldChar')); run._r[-1].set(qn('w:fldCharType'), 'separate')
        run2 = paragraph.add_run("[Z]")
        end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end'); run2._r.append(end)

    def add_bibl_field(doc):
        para = doc.add_paragraph(); run = para.add_run()
        for t in ['begin', 'instrText', 'separate']:
            el = OxmlElement('w:fldChar' if t != 'instrText' else 'w:instrText')
            if t == 'instrText': el.set(qn('xml:space'), 'preserve'); el.text = ' ADDIN ZOTERO_BIBL {"filename":"","custom":[]}'
            else: el.set(qn('w:fldCharType'), t)
            run._r.append(el)
        run2 = para.add_run("References will be generated by Zotero. Click Refresh.")
        end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end'); run2._r.append(end)

    def process_inline_text(paragraph, text, key2uri, key2info):
        MATH_SEP = '\uFFF0'
        def write_math(t, para, math_pl):
            parts = re.split(r'(\uFFF0(MATH|DMATH)\d+\uFFF0)', t)
            for part in parts:
                if part.startswith(MATH_SEP + 'MATH'):
                    idx = int(re.search(r'\d+', part).group()); latex_str = math_pl[idx][1]
                    run = para.add_run(' \\( ' + latex_str + ' \\) '); run.font.name = 'Consolas'; run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                elif part.startswith(MATH_SEP + 'DMATH'):
                    idx = int(re.search(r'\d+', part).group()); latex_str = math_pl[idx][1]
                    run = para.add_run(' [LaTeX: ' + latex_str + '] '); run.font.name = 'Consolas'; run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x33, 0x33, 0x33); run.italic = True
                elif t.strip(): para.add_run(part)

        def process_segment(seg, para):
            math_placeholders = []
            def save_math(m): math_placeholders.append(('inline', m.group(1).strip())); return MATH_SEP + 'MATH' + str(len(math_placeholders) - 1) + MATH_SEP
            def save_dmath(m): math_placeholders.append(('display', m.group(1).strip())); return MATH_SEP + 'DMATH' + str(len(math_placeholders) - 1) + MATH_SEP
            BS2 = chr(92) + chr(92)
            seg = re.sub(BS2 + r'\((.+?)' + BS2 + r'\)', save_math, seg)
            seg = re.sub(BS2 + r'\[(.+?)' + BS2 + r'\]', save_dmath, seg)
            seg = re.sub(r'\$\$(.+?)\$\$', save_dmath, seg)
            seg = re.sub(r'\$(.+?)\$', save_math, seg)
            bold_pattern = re.compile(r'\*\*(.+?)\*\*')
            pos = 0
            for m in bold_pattern.finditer(seg):
                pre = seg[pos:m.start()]
                if pre: write_math(pre, para, math_placeholders)
                run_b = para.add_run(m.group(1)); run_b.bold = True; pos = m.end()
            if pos < len(seg): write_math(seg[pos:], para, math_placeholders)

        parts = re.split(r'\[@([^\]]+)\]', text)
        for i, part in enumerate(parts):
            if i % 2 == 0:
                if part: process_segment(part, paragraph)
            else:
                keys = [k.strip() for k in part.replace('@', '').split(';')]; keys = [k for k in keys if k]
                if keys:
                    fj = make_citation_json(keys, key2uri, key2info)
                    if fj: add_field_code(paragraph, fj); paragraph.add_run(" ")

    def is_table_separator(line): return bool(re.match(r'^\|[\s\-:|\s]+\|$', line.strip()))
    def detect_table_lines(lines, start_idx):
        if start_idx >= len(lines): return None
        if not lines[start_idx].strip().startswith('|') or not lines[start_idx].strip().endswith('|'): return None
        header = [c.strip() for c in lines[start_idx].strip('|').split('|')]
        if start_idx + 1 >= len(lines): return None
        if not is_table_separator(lines[start_idx + 1]): return None
        data_rows = []; end_idx = start_idx + 2
        while end_idx < len(lines) and lines[end_idx].strip().startswith('|') and lines[end_idx].strip().endswith('|'):
            data_rows.append([c.strip() for c in lines[end_idx].strip('|').split('|')]); end_idx += 1
        return {'header': header, 'rows': data_rows, 'end': end_idx}

    log_cb("[1/4] 加载 Zotero 文献映射..."); key2uri, key2info = load_bbt()
    log_cb(f"  已加载 {len(key2uri)} 个 citation key → URI 映射")
    log_cb("[2/4] 读取 Markdown 源文件...")
    with open(md_path, "r", encoding="utf-8") as f: md_content = f.read()
    lines = md_content.split('\n'); log_cb(f"  共 {len(lines)} 行")
    log_cb("[3/4] 构建 DOCX 文档..."); doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(11)
    i = 0
    while i < len(lines):
        line = lines[i]; stripped = line.strip()
        if stripped == '## References' and i < 105:
            while i < len(lines) and not any(lines[i].strip().startswith(f'## {j}.') for j in range(3, 9)): i += 1
            continue
        if stripped == '## References' and i > 500: i = len(lines); continue
        if stripped == '\\[':
            i += 1
            while i < len(lines) and lines[i].strip() != '\\]': i += 1
            i += 1; continue
        heading_match = re.match(r'^(#{1,6})\s+(.+)', stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4); doc.add_heading(heading_match.group(2).strip(), level=level); i += 1; continue
        if stripped == '---': doc.add_paragraph(); i += 1; continue
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)', stripped)
        if img_match:
            img_rel_path = img_match.group(2); img_abs_path = Path(md_path).parent / img_rel_path
            caption_text = ""
            if i + 1 < len(lines) and lines[i+1].strip().startswith('*Figure'): caption_text = lines[i+1].strip().strip('*'); i += 1
            if img_abs_path.exists():
                try:
                    para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(); run.add_picture(str(img_abs_path), width=Inches(5.5))
                except Exception as e: log_cb(f"  [IMG ERROR] {img_rel_path}: {e}")
            i += 1; continue
        table_info = detect_table_lines(lines, i)
        if table_info:
            header = table_info['header']; rows = table_info['rows']
            num_cols = max(len(header), max((len(r) for r in rows), default=0))
            table = doc.add_table(rows=len(rows)+1, cols=num_cols); table.style = 'Light Grid Accent 1'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for col_idx, cell_text in enumerate(header):
                if col_idx < num_cols: cell = table.rows[0].cells[col_idx]; cell.text = ""; p = cell.paragraphs[0]; p.add_run(cell_text).bold = True
            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < num_cols: cell = table.rows[row_idx+1].cells[col_idx]; cell.text = ""; p2 = cell.paragraphs[0]; p2.add_run(cell_text)
            log_cb(f"  [TABLE] {num_cols} cols x {len(rows)} rows"); i = table_info['end']; continue
        if not stripped: i += 1; continue
        para = doc.add_paragraph(); process_inline_text(para, stripped, key2uri, key2info); i += 1
    log_cb("[4/4] 添加 Zotero 参考文献域..."); doc.add_paragraph(); add_bibl_field(doc)
    date_str = datetime.now().strftime("%y%m%d_%H%M%S")
    output_path = Path(output_dir) / f"output_zotero_{date_str}.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True); doc.save(str(output_path))
    import zipfile
    with zipfile.ZipFile(output_path) as z:
        xml = z.read('word/document.xml').decode('utf-8'); cit_count = len(re.findall(r'ADDIN ZOTERO_ITEM', xml))
    log_cb(f"\n✓ 转换完成!"); log_cb(f"  输出文件: {output_path}"); log_cb(f"  Zotero 引用域代码: {cit_count} 个")
    return str(output_path)

# ═══════════════════════════════════════════════════════════
# 5. FastAPI Application
# ═══════════════════════════════════════════════════════════

app = FastAPI(title="MinerU Desktop v3.0.4 Backend", version="3.0.4")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

async def broadcast_log(msg: str):
    for ws in list(state["log_clients"]):
        try: await ws.send_text(msg)
        except Exception: state["log_clients"].discard(ws)

def sync_log(msg: str):
    if loop: asyncio.run_coroutine_threadsafe(broadcast_log(msg), loop)

loop = None

@app.on_event("startup")
async def startup_event():
    global loop; loop = asyncio.get_event_loop()

class ParseRequest(BaseModel):
    citekeys: list; output_dir: str; backend: str = "vlm-auto-engine"; lang: str = "en"

class DirectFilesRequest(BaseModel):
    file_paths: list; output_dir: str; backend: str = "vlm-auto-engine"; lang: str = "en"

class DocxRequest(BaseModel):
    md_path: str; json_path: str; output_dir: str = ""

class ConfigData(BaseModel):
    mineru_path: str = ""; output_dir: str = ""; gpu_temp_threshold: int = 70
    backend: str = "vlm-auto-engine"; lang: str = "en"; theme: str = "light"
    mineru_api_token: str = ""; mineru_api_server: str = ""

@app.get("/api/health")
async def health():
    return {"status": "ok", "version": "3.2.2"}

@app.get("/api/quick")
async def quick_status():
    return {
        "status": "ok", "version": "3.2.1",
        "gpu": detect_gpu_fast(), "models": detect_models_fast(),
        "conda": cached_find_conda(), "mineru": cached_detect_mineru_env(),
        "parse_running": state["parse_running"], "setup_running": state["setup_running"],
    }

@app.get("/api/capability")
async def capability():
    return capability_assessment()

@app.get("/api/gpu/deep")
async def gpu_deep():
    return detect_gpu_deep()

@app.get("/api/status")
async def get_status():
    env = cached_detect_mineru_env(); gpu = cached_detect_gpu()
    models = cached_detect_models(); config = load_config()
    return {
        "mineru_installed": env is not None, "mineru_env": env,
        "conda_path": cached_find_conda(), "gpu": gpu, "models": models, "config": config,
        "parse_running": state["parse_running"], "parse_paused": state["parse_paused"],
        "parse_progress": state["parse_progress"], "setup_running": state["setup_running"],
        "docx_running": state["docx_running"],
    }

@app.get("/api/gpu")
async def get_gpu(): return cached_detect_gpu()

@app.get("/api/models")
async def get_models(): return cached_detect_models()

@app.get("/api/env")
async def get_env():
    return {
        "conda": cached_find_conda(), "mineru": cached_detect_mineru_env(),
        "gpu": cached_detect_gpu(), "ram": detect_system_ram(),
    }

# ── Memory Cleanup ──
@app.post("/api/cleanup")
async def api_cleanup():
    return cleanup_memory()

# ── MinerU Cloud API ──
@app.post("/api/api/validate")
async def validate_api_token(data: dict):
    token = data.get("token", "").strip()
    if not token: return {"valid": False, "message": "请输入 API Token"}
    try:
        req = urllib.request.Request("https://mineru.net/api/v4/api-usage",
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"})
        resp = urllib.request.urlopen(req, timeout=10)
        body = json.loads(resp.read().decode())
        return {"valid": True, "message": "Token 有效", "data": body}
    except urllib.error.HTTPError as e:
        code = e.code
        if code == 401: return {"valid": False, "message": "Token 无效"}
        if code == 403: return {"valid": False, "message": "Token 已过期或无权限"}
        return {"valid": False, "message": f"API 错误 (HTTP {code})"}
    except urllib.error.URLError as e:
        return {"valid": False, "message": "无法连接 MinerU 服务，请检查网络"}
    except json.JSONDecodeError:
        return {"valid": False, "message": "API 返回格式异常"}
    except Exception as e:
        return {"valid": False, "message": f"验证失败: {str(e)[:60]}"}

# ── Setup Wizard ──
@app.get("/api/setup/steps")
async def get_setup_steps():
    steps = []
    for s in SETUP_STEPS_DEF:
        try:
            skip = s.get("skip_if")
            if skip and skip():
                steps.append({"id": s["id"], "name": s["name"], "description": s["description"],
                    "location": s["location"], "size": s["size"], "action_text": s["action_text"],
                    "status": "skipped", "skip_reason": "硬件不支持，自动跳过"})
                continue
            done = s["check_done"]()
        except: done = False
        steps.append({"id": s["id"], "name": s["name"], "description": s["description"],
            "location": s["location"], "size": s["size"], "action_text": s["action_text"],
            "status": "done" if done else "pending"})
    return steps

@app.post("/api/setup/step/{step_id}")
async def execute_setup_step(step_id: str):
    state["setup_running"] = True
    def run():
        runner = SetupRunner(log_cb=lambda msg: sync_log(msg), progress_cb=lambda v: None)
        runner.execute_step(step_id)
        state["setup_running"] = False
        valid = any(s["id"] == step_id for s in SETUP_STEPS_DEF)
        if valid:
            step_info = next(s for s in SETUP_STEPS_DEF if s["id"] == step_id)
            done = step_info["check_done"]()
            sync_log(f"\n{'✓' if done else '✗'} {step_info['name']} {'完成' if done else '失败，请查看上方日志'}")
    executor.submit(run)
    return {"message": f"Step {step_id} started"}

@app.get("/api/setup/status")
async def setup_status():
    return {"running": state["setup_running"]}

# ── Zotero JSON Import ──
@app.post("/api/zotero/import")
async def import_zotero_json(data: dict):
    file_path = data.get("file_path", "")
    if not file_path or not Path(file_path).exists(): raise HTTPException(400, "Invalid file path")
    refs = parse_zotero_json(file_path)
    state["imported_references"] = refs
    config = load_config()
    output_dir = config.get("output_dir", str(Path.home() / "Desktop" / "mineru_output"))
    for ref in refs:
        md_output = Path(output_dir) / ref["citekey"] / f"{ref['citekey']}.md"
        ref["parsed"] = md_output.exists() and md_output.stat().st_size > 500
        ref["output_dir"] = output_dir
    sync_log(f"✓ 导入 {len(refs)} 条文献")
    return {"count": len(refs), "references": refs}

@app.get("/api/zotero/references")
async def get_references():
    return state["imported_references"]

# ── Direct File Import ──
@app.post("/api/files/import")
async def import_direct_files(data: dict):
    paths = data.get("file_paths", [])
    files = []
    for p in paths:
        path = Path(p)
        if not path.exists(): continue
        name = path.name; ext = path.suffix.lower()
        accept = ext in ['.pdf', '.docx', '.pptx', '.xlsx', '.png', '.jpg', '.jpeg', '.bmp']
        files.append({"name": name, "path": str(path), "size": path.stat().st_size, "type": ext, "valid": accept, "parsed": False})
    state["direct_files"] = files
    sync_log(f"✓ 导入 {len(files)} 个文件 ({sum(1 for f in files if f['valid'])} 个有效)")
    return {"count": len(files), "files": files}

@app.get("/api/files/list")
async def list_direct_files():
    return state["direct_files"]

@app.post("/api/files/parse")
async def start_direct_parse(data: dict):
    if state["parse_running"]: raise HTTPException(400, "Parse already running")
    file_paths = data.get("file_paths", [])
    output_dir = data.get("output_dir", "")
    backend = data.get("backend", "vlm-auto-engine")
    lang = data.get("lang", "en")
    parse_mode = data.get("parse_mode", "local")  # "local" | "cloud_api" | "cloud_model"
    api_token = data.get("api_token", "")

    if not file_paths: raise HTTPException(400, "No files to parse")
    if not output_dir: raise HTTPException(400, "No output directory")

    items = []
    for p in file_paths:
        path = Path(p)
        if not path.exists(): continue
        name = path.stem
        items.append({"name": name, "path": str(path), "pdf_path": str(path)})

    if not items: raise HTTPException(400, "No valid files found")
    state["parse_running"] = True; state["parse_cancel_requested"] = False; state["parse_paused"] = False
    state["parse_progress"] = {"current": 0, "total": len(items), "current_file": "", "file_progress": 0}

    def run_direct_parse():
        try:
            if parse_mode == "cloud_api" or parse_mode == "cloud_model":
                # 云端解析走 MinerU 官方 API
                _cloud_parse_core(items, output_dir, parse_mode, api_token, lang)
            else:
                # 本地解析
                env = cached_detect_mineru_env()
                if not env:
                    sync_log("✗ MinerU 本地环境未安装，无法本地解析")
                else:
                    _batch_parse_core(items, env, output_dir, backend, lang)
        finally:
            state["parse_running"] = False; state["parse_paused"] = False
            state["_current_proc"] = None; state["_current_stop"] = None

    executor.submit(run_direct_parse)
    return {"message": "Parse started (direct files)", "total": len(items)}

# ── Cloud Parse (MinerU 官方 API) ──
def _cloud_parse_core(items, output_dir, parse_mode, api_token, lang):
    """parse_mode: 'cloud_api' (官方在线API) 或 'cloud_model' (云端模型，付费更快)"""
    if not api_token:
        sync_log("✗ 未配置 MinerU Cloud API Token，请在设置中填入")
        return

    api_url = "https://mineru.net/api/v4"
    ok = 0; total = len(items)
    state["_current_stop"] = threading.Event()

    for idx, it in enumerate(items):
        if state["parse_cancel_requested"]:
            sync_log(f"⏹ 取消云端解析 (剩余 {total - idx} 篇)"); break

        sync_log(f"\n☁ [{idx+1}/{total}] {it['name']} (云端{'API' if parse_mode == 'cloud_api' else '模型'})")
        state["parse_progress"]["current"] = idx
        state["parse_progress"]["current_file"] = it["name"]
        state["parse_progress"]["file_progress"] = 10

        try:
            # 1. 申请上传链接
            req_body = json.dumps({
                "enable_formula": True, "enable_table": True,
                "language": lang if lang in ("en","ch") else "en",
                "is_ocr": True, "model_version": "vlm" if parse_mode == "cloud_model" else "pipeline"
            }).encode("utf-8")
            req = urllib.request.Request(
                f"{api_url}/file-urls/batch",
                data=req_body, method="POST",
                headers={"Authorization": f"Bearer {api_token}", "Content-Type": "application/json"}
            )
            with urllib.request.urlopen(req, timeout=30) as resp:
                rdata = json.loads(resp.read().decode("utf-8"))
            if rdata.get("code") != 0:
                sync_log(f"  ✗ 申请上传链接失败: {rdata.get('msg', '未知错误')}")
                continue
            batch_id = rdata["data"]["batch_id"]
            upload_urls = rdata["data"]["file_urls"]
            if not upload_urls:
                sync_log(f"  ✗ 申请上传链接失败: 返回为空")
                continue
            upload_url = upload_urls[0]

            # 2. 上传文件
            state["parse_progress"]["file_progress"] = 30
            with open(it["pdf_path"], "rb") as f:
                file_data = f.read()
            upload_req = urllib.request.Request(upload_url, data=file_data, method="PUT",
                                                headers={"Content-Type": "application/pdf"})
            urllib.request.urlopen(upload_req, timeout=300).read()
            sync_log(f"  ✓ 上传成功: {it['name']}")

            # 3. 轮询查询结果
            state["parse_progress"]["file_progress"] = 50
            max_wait = 600  # 10 分钟
            polled = 0
            zip_url = None
            while polled < max_wait:
                if state["parse_cancel_requested"]: break
                time.sleep(5); polled += 5
                qurl = f"{api_url}/extract-results/batch/{batch_id}"
                qreq = urllib.request.Request(qurl, headers={"Authorization": f"Bearer {api_token}"})
                with urllib.request.urlopen(qreq, timeout=30) as qresp:
                    qdata = json.loads(qresp.read().decode("utf-8"))
                extract = qdata.get("data", {}).get("extract_result", [])
                if extract:
                    er = extract[0]
                    state_msg = er.get("state", "")
                    if state_msg == "done":
                        zip_url = er.get("full_md_link") or er.get("zip_url")
                        break
                    elif state_msg == "failed":
                        sync_log(f"  ✗ 云端解析失败: {er.get('err_msg', '未知')}")
                        break
                    else:
                        state["parse_progress"]["file_progress"] = 50 + min(40, polled * 40 // max_wait)
                        continue
            if not zip_url:
                sync_log(f"  ⚠ {it['name']} 云端超时或失败")
                continue

            # 4. 下载并解压
            state["parse_progress"]["file_progress"] = 95
            with urllib.request.urlopen(zip_url, timeout=300) as zresp:
                zip_bytes = zresp.read()
            dst = Path(output_dir) / it["name"]
            dst.mkdir(parents=True, exist_ok=True)
            with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
                zf.extractall(dst)
            ok += 1
            state["parse_progress"]["file_progress"] = 100
            sync_log(f"  ✓ {it['name']} 云端解析完成")
        except Exception as e:
            sync_log(f"  ✗ {it['name']}: {e}")

    state["parse_progress"]["current"] = total
    sync_log(f"\n{'='*50}\n☁ 云端完成: {ok}/{total} 成功")

# ── Shared Batch Parse Core ──
def _batch_parse_core(items, env, output_dir, backend, lang):
    import tempfile as tmpfile_mod

    ok = 0; total = len(items)
    total_batches = (total + BATCH_SIZE - 1) // BATCH_SIZE

    for batch_idx in range(total_batches):
        if state["parse_cancel_requested"]:
            sync_log(f"⏹ 取消解析 (剩余 {total - batch_idx * BATCH_SIZE} 篇)"); break

        batch = items[batch_idx * BATCH_SIZE: (batch_idx + 1) * BATCH_SIZE]
        batch_num = batch_idx + 1
        names_batch = [it["name"] for it in batch]

        temp = get_gpu_temperature()
        if temp > GPU_TEMP_PAUSE_THRESHOLD:
            state["parse_paused"] = True
            state["parse_pause_reason"] = f"GPU 温度过高 ({temp}°C > {GPU_TEMP_PAUSE_THRESHOLD}°C)，暂停 {GPU_TEMP_PAUSE_SECONDS} 秒散热..."
            sync_log(state["parse_pause_reason"]); time.sleep(GPU_TEMP_PAUSE_SECONDS)
            state["parse_paused"] = False; state["parse_pause_reason"] = ""

        sync_log(f"\n{'─'*40}\n📦 批次 {batch_num}/{total_batches} | {', '.join(names_batch[:3])}{'...' if len(names_batch) > 3 else ''}")

        batch_dir = Path(tmpfile_mod.mkdtemp(prefix='mineru_batch_'))
        batch_out = batch_dir / '_output'
        batch_out.mkdir()

        for it in batch:
            try:
                shutil.copy2(it["pdf_path"], batch_dir / f"{it['name']}.pdf")
            except Exception as e:
                sync_log(f"  ✗ {it['name']}: 复制PDF失败 ({e})")

        retry = 0; batch_success = False
        batch_timeout = BATCH_SIZE * 720

        while retry < MAX_RETRIES and not batch_success:
            if retry > 0:
                sync_log(f"  🔄 批次 {batch_num} 第 {retry+1}/{MAX_RETRIES} 次重试 ({backend})...")
                time.sleep(3)
            cmd = [env["mineru"], "-p", str(batch_dir), "-o", str(batch_out), "-b", backend, "-l", lang]
            try:
                proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                                        text=True, encoding="utf-8", errors="replace")
                state["_current_proc"] = proc
                stop_reading = threading.Event()
                state["_current_stop"] = stop_reading
                def _read_stdout():
                    try:
                        for line in iter(proc.stdout.readline, ''):
                            if stop_reading.is_set(): break
                            if state["parse_cancel_requested"]: proc.terminate(); break
                            line = line.strip()
                            if line: sync_log(line)
                            m = re.search(r"Predict:\s+(\d+)%", line)
                            if m: state["parse_progress"]["file_progress"] = int(m.group(1))
                    except Exception: pass
                reader = threading.Thread(target=_read_stdout, daemon=True)
                reader.start()
                try:
                    proc.wait(timeout=batch_timeout)
                except subprocess.TimeoutExpired:
                    sync_log(f"  ⚠ 批次 {batch_num} 超时，终止中...")
                    try: proc.kill()
                    except: pass
                    retry += 1
                    stop_reading.set(); reader.join(timeout=5)
                    continue
                stop_reading.set()

                if proc.returncode == 0:
                    for it in batch:
                        src = batch_out / it["name"]
                        dst = Path(output_dir) / it["name"]
                        if src.exists():
                            if dst.exists(): shutil.rmtree(dst, ignore_errors=True)
                            shutil.move(str(src), str(dst))
                            ok += 1; state["parse_history"][it["name"]] = True
                            sync_log(f"  ✓ {it['name']}")
                        else:
                            state["parse_history"][it["name"]] = False
                            sync_log(f"  ✗ {it['name']} (minerU 未生成输出)")
                    batch_success = True
                else:
                    retry += 1
                    sync_log(f"  ✗ 批次 {batch_num} {backend} (exit={proc.returncode}){'，将重试...' if retry < MAX_RETRIES else ''}")

            except Exception as e:
                retry += 1
                sync_log(f"  ✗ 批次 {batch_num}: {e}{'，将重试...' if retry < MAX_RETRIES else ''}")

        try: shutil.rmtree(batch_dir, ignore_errors=True)
        except: pass

        state["parse_progress"]["current"] = min((batch_idx + 1) * BATCH_SIZE, total)

    sync_log(f"\n{'='*50}\n完成: {ok}/{total} 成功")

# ── Workflow A: Batch Parsing ──
@app.post("/api/parse/start")
async def start_parse(req: ParseRequest):
    if state["parse_running"]: raise HTTPException(400, "Parse already running")
    env = cached_detect_mineru_env()
    if not env: raise HTTPException(400, "MinerU not installed")
    refs = state["imported_references"]; ref_map = {r["citekey"]: r for r in refs}
    selected = [ref_map.get(ck) for ck in req.citekeys if ck in ref_map and ref_map[ck].get("pdf_exists")]
    if not selected: raise HTTPException(400, "No valid PDFs to parse")
    state["parse_running"] = True; state["parse_cancel_requested"] = False; state["parse_paused"] = False
    state["parse_progress"] = {"current": 0, "total": len(selected), "current_file": "", "file_progress": 0}

    items = [{"name": r["citekey"], "pdf_path": r["pdf_path"]} for r in selected]

    def run_parse():
        _batch_parse_core(items, env, req.output_dir, req.backend, req.lang)
        state["parse_running"] = False; state["parse_paused"] = False
        state["_current_proc"] = None; state["_current_stop"] = None

    executor.submit(run_parse)
    return {"message": "Parse started", "total": len(selected)}

@app.post("/api/parse/cancel")
async def cancel_parse():
    state["parse_cancel_requested"] = True
    proc = state.get("_current_proc")
    stop = state.get("_current_stop")
    if stop: stop.set()
    if proc:
        try: proc.terminate()
        except: pass
        try: proc.wait(timeout=2)
        except: pass
        try: proc.kill()
        except: pass
    state["parse_running"] = False
    state["parse_paused"] = False
    state["_current_proc"] = None
    state["_current_stop"] = None
    executor.submit(cleanup_memory)
    sync_log("⏹ 解析已取消，正在清理显存...")
    return {"message": "已取消，内存释放中"}

@app.get("/api/parse/status")
async def parse_status():
    return {
        "running": state["parse_running"],
        "paused": state["parse_paused"],
        "pause_reason": state["parse_pause_reason"],
        "progress": state["parse_progress"],
    }

# ── Workflow B: DOCX Generation ──
@app.post("/api/docx/generate")
async def generate_docx(req: DocxRequest):
    if state["docx_running"]: raise HTTPException(400, "DOCX generation already running")
    state["docx_running"] = True
    output_dir = req.output_dir or str(Path(req.md_path).parent)

    def run_docx():
        try:
            build_docx_with_zotero(req.md_path, req.json_path, output_dir, sync_log)
        except Exception as e:
            sync_log(f"✗ DOCX 生成失败: {e}")
        state["docx_running"] = False

    executor.submit(run_docx)
    return {"message": "DOCX generation started"}

@app.get("/api/docx/status")
async def docx_status():
    return {"running": state["docx_running"]}

# ── Config ──
@app.get("/api/config")
async def get_config():
    return load_config()

@app.post("/api/config")
async def post_config(data: ConfigData):
    config = data.dict()
    save_config(config)
    sync_log("✓ 设置已保存")
    return config

# ── WebSocket ──
@app.websocket("/ws/logs")
async def ws_logs(ws: WebSocket):
    await ws.accept(); state["log_clients"].add(ws)
    try:
        while True: await ws.receive_text()
    except WebSocketDisconnect:
        state["log_clients"].discard(ws)
    except Exception:
        state["log_clients"].discard(ws)

# ── Main ──
if __name__ == "__main__":
    port = int(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_PORT
    uvicorn.run("server:app", host="127.0.0.1", port=port, reload=False, log_level="warning")